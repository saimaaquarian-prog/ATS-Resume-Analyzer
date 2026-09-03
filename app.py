import io
import json
import os
import re
from typing import Any

import streamlit as st
from docx import Document
from pypdf import PdfReader
from google import genai
from google.genai import types

MODEL_NAME = "gemini-3.5-flash"
MAX_RESUME_CHARS = 50000

st.set_page_config(page_title="Resume ATS Analyzer", page_icon="📄", layout="wide")


def get_api_key() -> str | None:
    try:
        key = st.secrets.get("GEMINI_API_KEY")
    except Exception:
        key = None
    return key or os.getenv("GEMINI_API_KEY")


def extract_pdf_text(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    pages = [(page.extract_text() or "") for page in reader.pages]
    return "\n".join(pages).strip()


def extract_docx_text(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts).strip()


def extract_resume_text(uploaded_file) -> str:
    data = uploaded_file.getvalue()
    suffix = uploaded_file.name.lower().rsplit(".", 1)[-1]
    if suffix == "pdf":
        return extract_pdf_text(data)
    if suffix == "docx":
        return extract_docx_text(data)
    raise ValueError("Unsupported file type. Please upload a PDF or DOCX file.")


def clean_json_text(text: str) -> str:
    text = text.strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text, flags=re.I)
        text = re.sub(r"\s*```$", "", text)
    start, end = text.find("{"), text.rfind("}")
    if start == -1 or end == -1 or end <= start:
        raise ValueError("Gemini did not return a valid JSON object.")
    return text[start : end + 1]


def validate_analysis(data: dict[str, Any]) -> dict[str, Any]:
    required = ["overall_score", "category_scores", "summary", "strengths", "issues", "improvements", "bullet_rewrites", "ats_tips"]
    missing = [key for key in required if key not in data]
    if missing:
        raise ValueError("AI response is missing required fields: " + ", ".join(missing))

    score = int(data["overall_score"])
    if not 0 <= score <= 100:
        raise ValueError("Overall score must be between 0 and 100.")

    categories = data["category_scores"]
    if not isinstance(categories, dict):
        raise ValueError("Category scores must be an object.")
    expected = ["formatting", "structure", "content", "keywords", "experience", "professionalism"]
    for key in expected:
        if key not in categories:
            raise ValueError(f"Missing category score: {key}")
        categories[key] = max(0, min(100, int(categories[key])))

    for key in ["strengths", "issues", "improvements", "ats_tips"]:
        if not isinstance(data[key], list):
            raise ValueError(f"{key} must be a list.")

    if not isinstance(data["bullet_rewrites"], list):
        raise ValueError("bullet_rewrites must be a list.")

    data["overall_score"] = score
    return data


def build_prompt(resume_text: str, job_description: str) -> str:
    job_context = job_description.strip() or "No job description was provided. Perform a general ATS-style analysis."
    return f"""You are an expert resume and ATS-compatibility reviewer.

Analyze the resume below and return ONLY valid JSON matching the schema provided.

Important rules:
- This is an AI-generated ATS-style estimate, NOT an official ATS score and never a guarantee of passing an ATS.
- Analyze only evidence present in the resume and optional job description.
- Never invent employers, dates, degrees, skills, metrics, achievements, certifications, or experience.
- Never fabricate measurable results. If a metric is missing, recommend adding one only if the candidate can truthfully provide it.
- Evaluate ATS readability, standard sections, formatting risk, content quality, keywords, experience/project impact, and professionalism.
- If no job description is supplied, do not invent job-specific missing keywords. Set missing_keywords to an empty list.
- If a job description is supplied, identify relevant keywords/skills that are missing from the resume, but do not claim the candidate has a skill that is not evidenced.
- Keep recommendations specific and actionable.
- Do not reproduce the entire resume.
- Return JSON only; no Markdown and no commentary outside JSON.

Required JSON schema:
{{
  "overall_score": 0,
  "category_scores": {{
    "formatting": 0,
    "structure": 0,
    "content": 0,
    "keywords": 0,
    "experience": 0,
    "professionalism": 0
  }},
  "summary": "",
  "strengths": [""],
  "issues": [""],
  "missing_keywords": [""],
  "improvements": [""],
  "bullet_rewrites": [
    {{"original": "", "improved": ""}}
  ],
  "ats_tips": [""]
}}

Optional job description:
{job_context}

Resume:
{resume_text[:MAX_RESUME_CHARS]}
"""


def analyze_resume(resume_text: str, job_description: str, api_key: str) -> dict[str, Any]:
    client = genai.Client(api_key=api_key)
    response = client.models.generate_content(
        model=MODEL_NAME,
        contents=build_prompt(resume_text, job_description),
        config=types.GenerateContentConfig(
            response_mime_type="application/json",
            temperature=0.2,
        ),
    )
    if not response.text:
        raise ValueError("Gemini returned an empty response.")
    return validate_analysis(json.loads(clean_json_text(response.text)))


def score_label(score: int) -> str:
    if score >= 85:
        return "Excellent"
    if score >= 70:
        return "Good"
    if score >= 50:
        return "Needs Improvement"
    return "Poor"


def display_results(result: dict[str, Any], has_job_description: bool) -> None:
    score = result["overall_score"]
    st.subheader("📊 ATS-Style Score")
    col1, col2 = st.columns([1, 2])
    with col1:
        st.metric("Estimated ATS Score", f"{score}/100", score_label(score))
    with col2:
        st.progress(score / 100)
        st.caption("This is an AI-generated estimate, not an official ATS score.")

    st.subheader("Category Scores")
    labels = {
        "formatting": "Formatting",
        "structure": "Structure",
        "content": "Content",
        "keywords": "Keywords",
        "experience": "Experience",
        "professionalism": "Professionalism",
    }
    cols = st.columns(3)
    for index, (key, label) in enumerate(labels.items()):
        with cols[index % 3]:
            st.metric(label, f"{result['category_scores'][key]}/100")

    st.subheader("🤖 AI Resume Insight")
    st.write(result["summary"])

    left, right = st.columns(2)
    with left:
        st.markdown("### ✅ Strengths")
        for item in result["strengths"]:
            st.markdown(f"- {item}")
    with right:
        st.markdown("### ⚠️ Issues Found")
        for item in result["issues"]:
            st.markdown(f"- {item}")

    st.markdown("### 🚀 Recommended Improvements")
    for index, item in enumerate(result["improvements"], 1):
        st.markdown(f"**{index}.** {item}")

    if result["bullet_rewrites"]:
        with st.expander("✍️ Suggested Bullet Improvements"):
            for item in result["bullet_rewrites"]:
                st.markdown(f"**Before:** {item.get('original', '')}")
                st.markdown(f"**After:** {item.get('improved', '')}")
                st.divider()

    if has_job_description:
        with st.expander("🎯 Job Description Keyword Match"):
            missing = result.get("missing_keywords", [])
            if missing:
                st.write("Relevant keywords/skills not clearly present in the resume:")
                for item in missing:
                    st.markdown(f"- {item}")
            else:
                st.success("No major missing keywords were identified from the supplied job description.")

    with st.expander("📌 ATS Tips"):
        for item in result["ats_tips"]:
            st.markdown(f"- {item}")


st.title("📄 Resume ATS Analyzer")
st.caption("Analyze your resume for ATS-style compatibility and get practical AI-powered improvements.")
st.info("Your uploaded resume is processed for analysis and is not permanently stored by this app.")

uploaded_file = st.file_uploader("Upload your resume", type=["pdf", "docx"], help="Supported formats: PDF and DOCX")
job_description = st.text_area(
    "Paste Job Description (Optional)",
    height=180,
    placeholder="Paste the job description here for targeted keyword and skill matching...",
)

if uploaded_file:
    st.success(f"Uploaded: {uploaded_file.name}")

if st.button("🔍 Analyze Resume", type="primary", use_container_width=True):
    if uploaded_file is None:
        st.warning("Please upload a PDF or DOCX resume first.")
        st.stop()

    api_key = get_api_key()
    if not api_key:
        st.error("Gemini API key is not configured. Add GEMINI_API_KEY to Streamlit Secrets or your environment variables.")
        st.stop()

    try:
        with st.spinner("Extracting resume text..."):
            resume_text = extract_resume_text(uploaded_file)
        if len(re.sub(r"\s+", "", resume_text)) < 100:
            st.error("We could not extract enough text from this file. Please upload a text-based PDF or DOCX resume.")
            st.stop()

        with st.spinner("Analyzing your resume with Gemini Flash..."):
            analysis = analyze_resume(resume_text, job_description, api_key)

        st.success("Analysis completed successfully.")
        display_results(analysis, bool(job_description.strip()))

    except Exception as exc:
        message = str(exc).lower()
        if "429" in message or "rate" in message and "limit" in message:
            user_message = "The AI service is temporarily rate-limited. Please try again later."
        elif "api key" in message or "authentication" in message or "permission" in message:
            user_message = "The Gemini API configuration appears invalid. Check your GEMINI_API_KEY."
        elif "pdf" in message:
            user_message = "The PDF could not be read. Please upload a valid text-based PDF."
        elif "docx" in message or "zip" in message:
            user_message = "The DOCX file could not be read. Please upload a valid DOCX document."
        else:
            user_message = "The analysis could not be completed. Please try again or check your configuration."
        st.error(user_message)
        st.caption(f"Details: {exc}")
